import os
from google.cloud import vision_v1
from google.cloud.vision_v1 import types
from docx import Document
import io
import fitz

# Set your Google Cloud Vision API credentials
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r'C:\ResourceEase\visionAPI\GoogleCloudVisionAPI.json'
output_folder = r'C:\Users\USER\Desktop\ResourceEase\re\blog\static\blog\output'
def extract_handwritten_text_and_save_as_docx(image_path, output_folder):
    client = vision_v1.ImageAnnotatorClient()
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()

    image = vision_v1.types.Image(content=content)
    response = client.document_text_detection(image=image)

    doc = Document()
    docText = response.full_text_annotation.text
    doc.add_paragraph(docText)

    docx_output_path = os.path.join(output_folder, f"{os.path.splitext(os.path.basename(image_path))[0]}.docx")
    doc.save(docx_output_path)
    print(f"Extracted text saved as a DOCX file: {docx_output_path}")
    return docx_output_path

def extract_pages_from_pdf(input_file, output_dir):
    pdf_document = fitz.open(input_file)
    extracted_images = []
    os.makedirs(output_dir, exist_ok=True)

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        image_list = page.get_pixmap()

        image_file = os.path.join(output_dir, f'page_{page_number + 1}.png')
        image_list.save(image_file, "png")
        extracted_images.append(image_file)

    return extracted_images

def combine_docs_to_single_docx(input_dir, output_file):
    doc = Document()
    docx_files = [filename for filename in os.listdir(input_dir) if filename.lower().endswith('.docx')]

    for filename in sorted(docx_files):
        file_path = os.path.join(input_dir, filename)
        with open(file_path, 'rb') as file:
            sub_doc = Document(file)
            for element in sub_doc.element.body:
                doc.element.body.append(element)

    doc.save(output_file)
    print(f"Combined DOCX file saved as: {output_file}")

def convert_uploaded_pdf_to_docx(uploaded_file_path, output_folder):
    # Extract pages from the uploaded PDF
    image_files = extract_pages_from_pdf(uploaded_file_path, output_folder)

    # Extract and save handwritten text from images as separate DOCX files
    extracted_docs = []
    for image_path in image_files:
        docx_file = extract_handwritten_text_and_save_as_docx(image_path, output_folder)
        extracted_docs.append(docx_file)

    # Combine individual DOCX files into one single DOCX file
    output_docx = os.path.join(output_folder, 'combined_output.docx')
    combine_docs_to_single_docx(output_folder, output_docx)
    return output_docx
